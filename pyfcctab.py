"""First attempt at Python code to handle FCC tables"""

import docx
import re
import numpy as np
import astropy.units as units

import pandas as pd
from IPython.display import display, HTML
from versions import Version

# --------------------------------------------------------------------- Services
class Service:
    """A service that can be allocated to a band"""
    def __init__(self, name, abbreviation=None, science=False, science_support=False):
        self.name = name.lower().strip()
        if abbreviation is None:
            self.abbreviation = self.name
        self.science = science
        self.science_support = science_support

services_list = [
#    Service(""),
    Service("AERONAUTICAL MOBILE"),
    Service("AERONAUTICAL MOBILE-SATELLITE"),
    Service("AERONAUTICAL RADIONAVIGATION"),
    Service("AMATEUR"),
    Service("AMATEUR-SATELLITE"),
    Service("BROADCASTING"),
    Service("BROADCASTING-SATELLITE"),
    Service("EARTH EXPLORATION-SATELLITE", abbreviation="EESS", science=True),
    Service("FIXED"),
    Service("FIXED-SATELLITE"),
    Service("INTER-SATELLITE"),
    Service("LAND MOBILE"),
    Service("MARITIME MOBILE"),
    Service("MARITIME MOBILE-SATELLITE"),
    Service("MARITIME RADIONAVIGATION"),
    Service("METEOROLOGICAL AIDS", science_support=True),
    Service("METEOROLOGICAL-SATELLITE", science_support=True),
    Service("MOBILE except aeronautical mobile"),
    Service("MOBILE"),
    Service("MOBILE-SATELLITE except aeronautical mobile-satellite"),
    Service("MOBILE-SATELLITE except maritime mobile-satellite"),
    Service("MOBILE-SATELLITE"),
    Service("RADIO ASTRONOMY", abbreviation="RAS", science=True),
    Service("RADIODETERMINATION-SATELLITE"),
    Service("RADIOLOCATION"),
    Service("RADIOLOCATION-SATELLITE"),
    Service("RADIONAVIGATION"),
    Service("RADIONAVIGATION-SATELLITE"),
    Service("SPACE OPERATION"),
    Service("SPACE RESEARCH", science_support=True),
    Service("STANDARD FREQUENCY AND TIME SIGNAL", abbreviation="Time", science_support=True),
    Service("STANDARD FREQUENCY AND TIME SIGNAL-SATELLITE", abbreviation="Time", science_support=True),
]
services = {service.name: service for service in services_list}

# --------------------------------------------------------------------- Bounds
class NotBoundsError(Exception):
    """Exception used to flag failed parse of Bounds"""
    pass

class Bounds:
    """A frequeny bounds"""
    def __init__(self, low, high):
        self.bounds = (low, high)
        self.bandwidth = self.bounds[1] - self.bounds[0]
        self.center = 0.5*(self.bounds[0] + self.bounds[1])

    def __str__(self):
        result = f"{self.bounds[0]}-{self.bounds[1]}"
        return result

    def __eq__(self, a):
        return self.bounds == a.bounds

    def __ne__(self, a):
        return not (self == a)

    @classmethod
    def parse(cls, text, unit):
        """Turn a string giving a frequency range into a bounds object"""
        re_float = r"[0-9]+(?:\.[0-9]+)?"
        re_bounds = f"^({re_float})-({re_float})$"
        match = re.match(re_bounds, text)
        if match is not None:
            return cls(
                float(match.group(1))*unit,
                float(match.group(2))*unit)
        # Perhaps this is the "below the bottom" case.
        re_bottom = f"^Below ({re_float}) \(Not Allocated\)$"
        match = re.match(re_bottom, text)
        if match is None:
            raise NotBoundsError(f"Not a valid range: {text}")
        return cls(
            0.0*unit, float(match.group(1))*unit)
    
# --------------------------------------------------------------------- Allocations
class Allocation:
    """An entry allocating a service to a band"""
    def __str__(self):
        """Return a string representation of an Allocation"""
        if self.primary:
            result = self.service.upper()
        else:
            result = self.service.capitalize()
        if len(self.modifiers) != 0:
            result += " " + " ".join([f"({m})" for m in self.modifiers])
        if len(self.footnotes) != 0:
            result += " " + " ".join(self.footnotes)
        return (result)

    def __eq__(self,a):
        if self.primary != a.primary:
            return False
        if self.modifiers != a.modifiers:
            return False
        if self.footnotes != a.footnotes:
            return False
        if self.co_primary != a.co_primary:
            return False
        if self.exclusive != a.exclusive:
            return False
        return True

    def __ne__(self, a):
        return not (self == a)

    @classmethod
    def parse(cls, line):
        """Take a complete line of text and turn into an Allocation"""
        # Work out which service this is.
        service = None
        ll = line.lower()
        for s in services.keys():
            if ll[0:len(s)] == s:
                if service is None:
                    service = s
                else:
                    if len(s) > len(service):
                        service = s
        # If not a service then quit
        if service is None:
            return None
        # Look at the remainder of the line
        invocation = line[0:len(service)]
        if len(invocation) > 0:
            first_word = invocation.split()[0]
        else:
            first_word = invocation
        primary = first_word.isupper()
        remainder = line[len(service):].strip()
        # Anyting in parentheses becomes a modifiers
        modifiers = []
        while len(remainder) > 0:
            if remainder[0] == r"(":
                modifier = remainder[1:remainder.index(r")")]
                modifiers.append(modifier)
                remainder = remainder[len(modifier)+2:].strip()
            else:
                break
        # Now the remainder (if anything) must be footnotes
        footnotes = remainder.split()
        # Create the result
        result = cls()
        result.service = service #.strip()
        result.modifiers = modifiers
        result.footnotes = footnotes
        result.primary = primary
        return result
            
# --------------------------------------------------------------------- Bands
def cell2text(cell, munge=False):
    result = [p.text for p in cell.paragraphs]
    if munge:
        return "\n".join(result)
    else:
        return result
    
def first_line(cell):
    return cell2text(cell)[0].strip()

def last_line(cell):
    return cell2text(cell)[-1].strip()

class NotBandError(Exception):
    """Exception used to indicate failed parse of band"""
    pass

class Band:
    """A frequency bounds and the allocations thereto (i.e., contents of a cell in the FCC tables)"""
    def __str__(self, separator="\n"):
        """Return a string representation of a Band"""
        result = str(self.bounds) + separator
        for a in self.allocations:
            result = result + str(a) + separator
        if self.footnotes != "":
            result = result + separator + " ".join(self.footnotes)
        return (result)

    def compact_str(self):
        return self.__str__(separator="/")

    def __eq__(self,a):
        """Compare two sets of Band information"""
        if self.bounds != a.bounds:
            return False
        if len(self.allocations) != len(a.allocations):
            return False
        for sa, aa in zip(self.allocations, a.allocations):
            if sa != aa:
                return False
        return self.footnotes == a.footnotes
        
    def __ne__(self, a):
        return not (self == a)

    def all_footnotes(self):
        result = self.footnotes
        for a in self.allocations:
            result = result + a.footnotes
        return list(set(result))

    def definitelyUSA(self):
        for f in self.all_footnotes():
            if f[0] != "5" and f[0] != "(":
                return True
        return False

    @classmethod
    def parse(cls, cell, unit):
        """Parse a table cell into a Band"""
        if cell is None:
            raise NotBandError("Cell is None")
        if type(cell) == docx.table._Cell:
            text = cell2text(cell)
        else:
            text = cell
        # Now work out which lines are continuations of other lines and string them together
        lines = []
        line = None
        for t in text:
            if len(t) != 0:
                t0 = t[0]
            else:
                t0 = ""
            is_continuation = (t0 == " ")
            if line is None:
                is_continuation = False
            else:
                if len(line) == 0:
                    is_continuation = False
            if is_continuation:
                if line[-1] != '-':
                    line = line + " " + t.strip()
                else:
                    line = line + t.strip()
            else:
                if line is not None:
                    lines.append(line)
                line = t.strip()
        if line is not None:
            lines.append(line.strip())
        # Now the first line should be a frequency range
        if len(lines) == 0:
            return None
        try:
            bounds = Bounds.parse(lines[0], unit)
        except NotBoundsError:
            raise NotBandError("Text doesn't start with bounds, so not a band")
        # Now the remainder will either be allocations, blanks or collections of footnotes
        service_names = services.keys()
        footnotes = []
        allocations = []
        primary_allocations = []
        secondary_allocations = []
        for l in lines[1:]:
            if l.strip() == "":
                continue
            # if footnotes is not None:
            #     raise ValueError("Gone past footnotes and cell not empty")
            # See if this line conveys an allocation
            allocation = Allocation.parse(l)
            if allocation is not None:
                if len(footnotes) != 0:
                    raise NotBandError("Back to allocation after footnotes")
                if allocation.primary:
                    primary_allocations.append(allocation)
                else:
                    secondary_allocations.append(allocation)
                allocations.append(allocation)
            else:
                footnotes += l.split()
        # Done looping over the lines, so tidy things up.
        if footnotes is None:
            footnotes = []
        for a in allocations:
            if a.primary:
                a.co_primary = len(primary_allocations) > 1
            else:
                a.co_primary = False
            a.exclusive = len(allocations) == 1
            a.bounds = bounds
        # Now create a Band to hold the result
        result = cls()
        result.bounds = bounds
        result.allocations = allocations
        result.primary_allocations = primary_allocations
        result.secondary_allocations = secondary_allocations
        result.footnotes = footnotes
        result._lines = lines
        result._text = text
        return result
    
# --------------------------------------------------------------------- Reading

def _dump_cells(cells):
    for i, c in enumerate(cells):
        print (f"-------------- Cell {c}")
        print (cell2text(c))

def pretty_print(df):
    return display(HTML(df.to_html().replace(r"\n","<br>")))

class OtherTable(Exception):
    pass
class FCCTableError(Exception):
    pass

def harvest_footnotes(fccfile, n_tables=64):
    """Go through the file and try to collect all the footnotes"""
    footnotes = []
    for table in fccfile.tables[0:n_tables]:
        for row in table.rows:
            for cell in row.cells:
                try:
                    band = Band.parse(cell, units.kHz)
                    footnotes += band.all_footnotes()
                except (ValueError):
                    pass
    return list(set(footnotes))

def _get_empty_collections():
    """Return a pair of lists each containing three empty lists"""
    itu_collections = []
    for i in range(3):
        itu_collections.append(list())
    fcc_collections = []
    for i in range(3):
        fcc_collections.append(list())
    return itu_collections, fcc_collections

    
def parse_table(table, unit, version=Version("20200818"), dump_raw=False, dump_ordered=False):
    """Go through one table in the document and return a guess as to its contents"""
    # Check that this is the kind of table we can cope with
    pass
    # First get the first row and work out if this is a table with headers or not
    header = first_line(table.rows[0].cells[0])
    n_rows = len(table.rows)
    header_prefix = "Table of Frequency Allocations"
    has_header = header[0:len(header_prefix)] == header_prefix
    page = None
    if has_header:
        page = last_line(table.rows[0].cells[-1])
        first_useful_row = 3
        # Get the units from the remainder of the header
        words = header[len(header_prefix):].split()
        unit = units.Unit(words[1])        
    else:
        first_useful_row = 0
    # Get the last row and check it's not just full of page numbers
    try:
        footer = first_line(table.rows[-1].cells[0])
    except IndexError:
        footer = ""
    footer_prefix = "Page"
    has_footer = footer[0:len(footer_prefix)] == footer_prefix
    if has_footer:
        last_useful_row = n_rows-2
    else:
        last_useful_row = n_rows-1
    # Also, get the page number if it's in the last row/column (or thereabouts)
    last_working_row = -1
    if not has_header:
        for ir in [-1,-2,-3]:
            try:
                page = last_line(table.rows[ir].cells[-1])
                last_working_row = ir
                break
            except IndexError:
                pass
    # Some page numbers are hard to deduce, these we patch
    page = version.patch_page(page)
    # Check to see if the table ends with a page number
    entry = first_line(table.rows[last_working_row].cells[-1])
    ends_with_page_number = entry[0:len(footer_prefix)] == footer_prefix

    # Possibly dump the raw table
    if dump_raw:
        frame = pd.DataFrame()
        for r in table.rows:
            entries = []
            for c in r.cells:
                try:
                    this_bottom = c._element.bottom
                except (ValueError):
                    this_bottom = None
                entries.append(
                    c.text +
                    f"<{c._element.left},{c._element.right}>, <{c._element.top},{this_bottom}>")
            frame = frame.append(pd.Series(entries), ignore_index=True)
        print(page)
        pretty_print(frame)

    n_cols_per_row = [len(r.cells) for r in table.rows]
    max_cols = max(n_cols_per_row)

    # Look at how each cell spans the grid (left and right information)
    left, right, top, bottom = [
        np.zeros(shape=[n_rows, max_cols], dtype=int)
        for i in range(4)]
    for ir, r in enumerate(table.rows):
        for ic, c in enumerate(r.cells):
            left[ir, ic] = c._element.left
            right[ir, ic] = c._element.right
            top[ir, ic] = c._element.top
            try:
                bottom[ir, ic] = c._element.bottom
            except ValueError:
                bottom[ir, ic] = top[ir, ic]+1
    # print(left)
    # print(right)
    # print(top)
    # print(bottom)
    # print(np.max(bottom), n_rows)
    assert np.max(bottom)==n_rows, "Confused about the number of rows"
    
    # Build up a new data structure for the cells in the proper order
    max_boxes = np.max(right)
    n_boxes_per_row = np.max(right, axis=1)
    ordered = []
    for ir in range(n_rows):
        boxes = [None]*max_boxes
        ordered.append(boxes)
    for rIn, r in enumerate(table.rows):
        for cIn, c in enumerate(r.cells):
            for rOut in range(top[rIn,cIn], bottom[rIn,cIn]):
                for cOut in range(left[rIn,cIn], right[rIn,cIn]):
                    new_value = cell2text(c)
                    current_value =  ordered[rOut][cOut]
                    if current_value != None and current_value != new_value:
                        raise FCCTableError(f"Trampled unexpectedly {rOut},{cOut} from {rIn},{cIn}")
                    else:
                        ordered[rOut][cOut] = new_value

    # Possibly dump the ordered table
    if dump_ordered:
        frame = pd.DataFrame(columns=range(max_boxes))
        for boxes in ordered:
            frame = frame.append(pd.Series(boxes), ignore_index=True)
        print(page)
        pretty_print(frame)

    # Now go through and cut out the header row
    ordered = ordered[first_useful_row:last_useful_row+1]
    left = left[first_useful_row:last_useful_row+1,:]
    right = right[first_useful_row:last_useful_row+1,:]
    top = top[first_useful_row:last_useful_row+1,:]
    bottom = bottom[first_useful_row:last_useful_row+1,:]
    n_rows = n_rows - first_useful_row

    # Create a pair of lists to hold the results
    itu_collections, fcc_collections = _get_empty_collections()
    for ir, boxes in enumerate(ordered):
        # OK, get the layout for this page/row
        layout = version.get_layout(page, ir+first_useful_row)
        assert layout[6] == "/", f"Bad layout: {layout}"
        if int(layout[7:]) != max_boxes:
            raise ValueError("Supplied layout does not match table")
        sources = [int(l,16) for l in layout[0:6]]
        # First the ITU entries
        entries = []
        for i in range(3):
            itu_collections[i].append(boxes[sources[i]])
        for i in range(3):
            fcc_collections[i].append(boxes[sources[i+3]])

    # Finish off
    diagnostics = {
        "page": page,
        "has_header": has_header}
    return itu_collections, fcc_collections, unit, diagnostics

class DigestError(Exception):
    """Error raised when collection can't be digested"""
    pass

def _digest_collection(entries, unit):
    """Go through a collection column, merge what needs to be merged and parse into bands"""
    output_collection = []
    accumulator = None
    previous = None
    for entry in entries:
        # See if this one is at least the start of a band
        try:
            if entry is not None:
                test = Band.parse(entry, unit)
            # OK, it worked (or it's not none), so dispatch the one
            # we've been working on thus far.  However, be careful,
            # there are circmstances in which this can simply be a
            # repeat of the previous one.
            if entry != previous and accumulator is not None:
                output_collection.append(Band.parse(accumulator, unit))
            accumulator = entry
        except NotBandError:
            # The current entry isn't a band, it must be finishing the
            # band in the accumulator.  Check it's not just a repeat.
            if entry != previous:
                if accumulator is None:
                    accumulator = entry
                else:
                    if entry is not None:
                        accumulator = accumulator + entry
        previous = entry
    # OK, now we've got to the end, deal with the final entry
    try:
        output_collection.append(Band.parse(accumulator, unit))
    except NotBandError:
        pass
        # raise DigestError("Unable to parse the final entry")
    return output_collection

def parse_file(fccfile, table_range=range(0,65), **kwargs):
    """Go through tables in fcc file and parse them"""

    version = Version("20200818")
    tables = fccfile.tables
    unit = units.dimensionless_unscaled
    # Setup empty result
    itu_accumulators, fcc_accumulators = _get_empty_collections()
    itu_collections, fcc_collections = _get_empty_collections()

    for it in table_range:
        table = tables[it]
        these_itu_entries, these_fcc_entries, unit, diagnostics = parse_table(
            table, unit, version, **kwargs)
        # If this table has got a header, then dispatch the previously
        # collected table entries.
        if diagnostics["has_header"]:
            for i in range(3):
                itu_collections[i] += _digest_collection(itu_accumulators[i], unit)
            itu_accumulators, fcc_accumulators = _get_empty_collections()
        # Now add the new entries to the accumulators
        for i in range(3):
            itu_accumulators[i] += these_itu_entries[i]
            fcc_accumulators[i] += these_fcc_entries[i]
    # Now displatch the final table
    for i in range(3):
        itu_collections[i] += _digest_collection(itu_accumulators[i], unit)
    # Code needed here to properly annotate the FCC collections.
    
    return itu_collections, fcc_collections
        
        
